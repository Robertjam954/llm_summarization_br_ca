"""
Generate a Word document summarising the V1 Prompt Validation EDA figures.
Output: reports/eda_v1_validation/V1_validation_EDA_figure_summary.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUT_DIR      = PROJECT_ROOT / "reports" / "eda_v1_validation"
DOCX_PATH    = OUT_DIR / "V1_validation_EDA_figure_summary.docx"
OUT_DIR.mkdir(parents=True, exist_ok=True)

doc = Document()

# ── Styles ─────────────────────────────────────────────────────────────────────
style_normal = doc.styles["Normal"]
style_normal.font.name  = "Calibri"
style_normal.font.size  = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.runs[0].font.color.rgb = RGBColor(0x00, 0x72, 0xB2)   # Okabe-Ito blue
    return h

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    return p

def add_field(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    p.add_run(value)
    return p

def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"⚠  {text}")
    run.italic  = True
    run.font.color.rgb = RGBColor(0xCC, 0x79, 0xA7)   # Okabe-Ito purple
    return p

# ── Title page ─────────────────────────────────────────────────────────────────
title = doc.add_heading("V1 Prompt Validation – EDA Figure Summary", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

intro = doc.add_paragraph(
    "This document summarises the six exploratory data analysis (EDA) figures "
    "generated from the V1 prompt validation dataset (n = 200 patients). "
    "Figures compare human-authored vs. AI-generated (V1 prompt) clinical breast "
    "radiology summary annotations against source documents across 14 clinical features. "
    "All output files are saved to: reports/eda_v1_validation/."
)
intro.paragraph_format.space_after = Pt(12)

doc.add_paragraph()

# ── Figure 1 ───────────────────────────────────────────────────────────────────
add_heading("Figure 1 – Feature Presence in Source Documents", level=1)
add_field("Output file", "fig1_feature_presence.png")
add_field("Plot type",   "Single bar chart (Okabe-Ito blue bars)")
add_body("Shows how frequently each of the 14 clinical features was present in the "
         "source radiology documents across the 200 validation cases.")
doc.add_paragraph("Key findings:", style="List Bullet").runs[0].bold = True
add_bullet("Lesion laterality (100%) and histologic diagnosis (100%) were present in every case.")
add_bullet("Lesion size (98.0%), lesion location (99.5%), chronology preservation (99.5%), "
           "and biopsy method (99.5%) were near-universally present.")
add_bullet("Receptor status (92.5%) and accurate clip placement (90.0%) were present "
           "in the large majority of cases.")
add_bullet("Extent of disease (44.0%) and additional MRI enhancement (28.5%) were "
           "present in fewer than half of source documents.")
doc.add_paragraph()

# ── Figure 2 ───────────────────────────────────────────────────────────────────
add_heading("Figure 2 – Accurate (%) Annotations: Human vs. AI", level=1)
add_field("Output file", "fig2_pct_correct.png")
add_field("Plot type",   "Side-by-side bar chart with 95% CI error bars and significance stars")
add_body("Compares the percentage of annotations correctly identified per feature "
         "(source = 1, annotation = 1) for human vs. AI. "
         "Error bars represent published 95% confidence intervals. "
         "Significance stars indicate features where human and AI accuracy differed "
         "significantly (McNemar's test; * p<0.05, ** p<0.01, *** p<0.001).")
doc.add_paragraph("Key findings:", style="List Bullet").runs[0].bold = True
add_bullet("AI outperformed humans on 12 of 14 features.")
add_bullet("The largest differences were for invasive component size (AI 93.8% vs. "
           "Human 30.1%, p<0.001), lymph node status (95.1% vs. 71.8%, p<0.001), "
           "and accurate clip placement (96.7% vs. 86.1%, p<0.001).")
add_bullet("Human outperformed AI only for receptor status (Human 96.2% vs. AI 85.9%, p=0.002).")
add_bullet("Significant features (p<0.05): Extent (*), Accurate Clip Placement (***), "
           "Workup Recommendation (*), Lymph Node (***), Biopsy Method (**), "
           "Invasive Component Size (***), Receptor Status (**).")
add_bullet("No significant difference for: Lesion Size, Laterality, Location, "
           "Calcifications/Asymmetry, MRI Enhancement, Chronology, Histologic Diagnosis.")
doc.add_paragraph()

# ── Figure 3 ───────────────────────────────────────────────────────────────────
add_heading("Figure 3 – Omitted (%) Annotations: Human vs. AI", level=1)
add_field("Output file", "fig3_pct_omitted.png")
add_field("Plot type",   "Side-by-side bar chart with 95% CI error bars")
add_body("Compares the percentage of features present in the source document that were "
         "omitted (annotation = 2) by human vs. AI annotators.")
doc.add_paragraph("Key findings:", style="List Bullet").runs[0].bold = True
add_bullet("Human annotators had markedly higher omission rates than AI on most features.")
add_bullet("Invasive component size had the highest human omission rate (69.9% vs. AI 4.8%), "
           "consistent with AI's superiority in extracting pathology data.")
add_bullet("Lymph node status: Human 28.2% omission vs. AI 4.9%.")
add_bullet("Workup recommendation: Human 16.8% vs. AI 6.8%; "
           "Additional MRI enhancement: Human 14.0% vs. AI 3.5%.")
add_bullet("AI had higher omission than humans only for receptor status (AI 13.0% vs. Human 3.2%).")
doc.add_paragraph()

# ── Figure 4 ───────────────────────────────────────────────────────────────────
add_heading("Figure 4 – Fabricated (%) Annotations: Human vs. AI", level=1)
add_field("Output file", "fig4_pct_fabricated.png")
add_field("Plot type",   "Side-by-side bar chart with 95% CI error bars")
add_body("Compares the percentage of features present in the source document that were "
         "fabricated (annotation = 3) — i.e., mentioned in the summary but contradicted "
         "or not supported by the source document.")
doc.add_paragraph("Key findings:", style="List Bullet").runs[0].bold = True
add_bullet("Overall fabrication rates were very low for both annotators across all features.")
add_bullet("Lesion size had the highest AI fabrication rate (2.0%) and human fabrication (1.0%).")
add_bullet("Invasive component size AI fabrication: 1.4%; "
           "Workup recommendation AI: 1.2%; Receptor status AI: 1.1%.")
add_bullet("8 of 14 features had 0% human fabrication; "
           "6 of 14 features had 0% AI fabrication.")
add_bullet("Wide 95% CIs for low-count features reflect small sample sizes "
           "within those feature subgroups.")
doc.add_paragraph()

# ── Figure 5 ───────────────────────────────────────────────────────────────────
add_heading("Figure 5 – Per-Feature AI Fabrication Rate: V1 vs. V2 Prompts", level=1)
add_field("Output file", "fig5_fab_v1_v2_per_feature.png")
add_field("Plot type",   "Side-by-side bar chart (V1 = Okabe-Ito blue; V2 = Okabe-Ito purple)")
add_body("Compares per-feature AI fabrication rates between the unstructured V1 prompt "
         "and the structured V2 prompt across all 14 clinical features. "
         "This figure is designed to assess whether the V2 prompt reduces fabrication "
         "relative to V1.")
doc.add_paragraph("Current status:", style="List Bullet").runs[0].bold = True
add_bullet("V1 (Unstructured) bars are populated from the V1 validation dataset (n = 200).")
add_bullet("V2 (Structured) bars are currently blank — V2 validation data not yet available.")
add_note(
    "ACTION REQUIRED: Once V2 validation data are available, update the variable "
    "'v2_fab_vals' in v1_prompt_validation_eda.py (line ~442) with the computed "
    "per-feature AI fabrication percentages. Re-running the script will automatically "
    "populate Figure 5 with the V2 bars."
)
doc.add_paragraph()

# ── Figure 6 ───────────────────────────────────────────────────────────────────
add_heading("Figure 6 – Overall AI Fabrication Rate: V1 vs. V2 Prompts", level=1)
add_field("Output file", "fig6_overall_fab_v1_v2.png")
add_field("Plot type",   "Two-bar chart (V1 = Okabe-Ito blue; V2 = Okabe-Ito purple)")
add_body("Shows the overall AI fabrication rate for the V1 and V2 prompts, computed as "
         "the unweighted mean of per-feature AI fabrication rates across all 14 features.")
doc.add_paragraph("Current status:", style="List Bullet").runs[0].bold = True
add_bullet("V1 overall mean AI fabrication rate = 0.64% (mean across 14 features).")
add_bullet("V2 bar is blank — V2 validation data not yet available.")
add_note(
    "ACTION REQUIRED: Once V2 validation data are available, update the variable "
    "'v2_overall_fab' in v1_prompt_validation_eda.py (line ~491) with the computed "
    "overall V2 AI fabrication rate. Re-running the script will populate Figure 6."
)
doc.add_paragraph()

# ── Closing note ───────────────────────────────────────────────────────────────
add_heading("Notes on V2 Data Integration", level=1)
add_body(
    "Figures 5 and 6 are intentionally structured as V1 vs. V2 comparison figures but "
    "currently display only V1 data. The following steps should be taken once V2 "
    "validation data are available:"
)
add_bullet(
    "Compute per-feature AI fabrication rates from the V2 validation datasheet "
    "using the same pct_metric() logic in v1_prompt_validation_eda.py."
)
add_bullet(
    "Replace the 'v2_fab_vals' placeholder (np.full array of NaN, ~line 442) with "
    "the 14-element array of V2 per-feature AI fabrication percentages."
)
add_bullet(
    "Replace the 'v2_overall_fab' placeholder (np.nan, ~line 491) with the overall "
    "V2 mean fabrication rate."
)
add_bullet(
    "Re-run v1_prompt_validation_eda.py. Figures 5 and 6 will automatically update "
    "with the V2 bars in Okabe-Ito purple."
)
add_body(
    "No other script changes are required to display V2 data in these figures."
)

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(DOCX_PATH)
print(f"Saved: {DOCX_PATH}")
