"""
Phase 2: Document Similarity Analysis
=====================================
Uses Voyage AI embeddings to compute semantic similarity between v1 and v2
summaries across 14 fabrication cases. Performs DBSCAN clustering and
generates t-SNE visualizations.

Usage:
    python fabrication_analysis/scripts/phase2_document_similarity.py

Requirements:
    pip install voyageai scikit-learn matplotlib seaborn
    Set VOYAGE_API_KEY in .env

Outputs (all under experiments/runs/v1_v2_comparison/similarity/):
    similar_case_pairs.json
    error_clusters.json
    similarity_map_tsne.png
    similarity_analysis_report.json
"""

from __future__ import annotations

import json
import os
import sys
from pathlib import Path

import time
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from dotenv import load_dotenv
from docx import Document

load_dotenv()

PROJECT_ROOT = Path(os.getenv("PROJECT_ROOT",
    r"C:\Users\jamesr4\OneDrive - Memorial Sloan Kettering Cancer Center"
    r"\Documents\GitHub\llm_summarization_br_ca"))
DATA_PRIVATE = Path(os.getenv("DATA_PRIVATE_DIR",
    r"C:\Users\jamesr4\loc\data_private"))
V2_DIR = Path(
    r"C:\Users\jamesr4\OneDrive - Memorial Sloan Kettering Cancer Center"
    r"\Documents\Research\Projects\moo\llm_summary\data\raw\fabrications_iteration_2"
)

FAB_XLSX     = DATA_PRIVATE / "raw" / "ai_fabrications_dataset.xlsx"
V1_PATHS_CSV = DATA_PRIVATE / "raw" / "v1_summary_paths.csv"
V2_PATHS_CSV = DATA_PRIVATE / "raw" / "v2_summary_paths.csv"
OUT_DIR      = PROJECT_ROOT / "experiments" / "runs" / "v1_v2_comparison" / "similarity"
REPORTS_DIR  = PROJECT_ROOT / "reports"

OUT_DIR.mkdir(parents=True, exist_ok=True)

VOYAGE_MODEL = "voyage-3"     # or "voyage-3-lite" for lower cost
TOP_K_PAIRS  = 10             # top similar pairs to report
DBSCAN_EPS   = 0.25           # cosine distance threshold for clustering
DBSCAN_MIN   = 2              # min samples per cluster

sys.path.insert(0, str(PROJECT_ROOT))


# ── Text extraction ────────────────────────────────────────────────────────────
def extract_docx_text(path: str) -> str:
    try:
        doc   = Document(str(path))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for trow in table.rows:
                cells = " | ".join(c.text.strip() for c in trow.cells if c.text.strip())
                if cells:
                    parts.append(cells)
        return "\n".join(parts).strip()
    except Exception as exc:
        return f"[ERROR: {exc}]"


def read_v2_text(path: str) -> str:
    try:
        return Path(path).read_text(encoding="utf-8", errors="ignore")[:1500]
    except Exception as exc:
        return f"[ERROR: {exc}]"


# ── Load data ──────────────────────────────────────────────────────────────────
def load_data() -> pd.DataFrame:
    df_fab  = pd.read_excel(FAB_XLSX)
    df_v1   = pd.read_csv(V1_PATHS_CSV).rename(columns={"summary_path": "v1_path"})
    df_v2   = pd.read_csv(V2_PATHS_CSV).rename(columns={"summary_path": "v2_path"})

    ai_cols = [c for c in df_fab.columns if c.endswith("_status_ai")]
    for c in ai_cols:
        df_fab[c] = pd.to_numeric(df_fab[c], errors="coerce")

    df_fab["surgeon_last"] = df_fab["surgeon"].str.split(",").str[0].str.strip()
    df_fab["fab_features"] = df_fab.apply(
        lambda row: [c.replace("_status_ai", "") for c in ai_cols if row[c] == 3],
        axis=1,
    )

    df = df_fab.merge(df_v1, on="mrn", how="left").merge(df_v2, on="mrn", how="left")
    df["v1_text"] = df["v1_path"].apply(
        lambda p: extract_docx_text(p) if pd.notna(p) else ""
    )
    df["v2_text"] = df["v2_path"].apply(
        lambda p: read_v2_text(p) if pd.notna(p) and Path(p).exists() else ""
    )
    return df


# ── Voyage AI embeddings ───────────────────────────────────────────────────────
def get_embeddings(texts: list[str], voyage_client,
                   batch_size: int = 3, rpm_limit: int = 3) -> np.ndarray:
    """
    Embed texts using Voyage AI with batched requests.
    Respects free-tier rate limit (3 RPM) via sleep between batches.
    """
    sleep_secs = 60.0 / rpm_limit + 2   # 22 s between requests at 3 RPM
    all_embs   = []
    batches    = [texts[i:i + batch_size] for i in range(0, len(texts), batch_size)]

    for idx, batch in enumerate(batches):
        # Trim each text to ~1500 chars (~375 tokens) to stay within 10K TPM
        trimmed = [t[:1500] for t in batch]
        est_tokens = sum(len(t) // 4 for t in trimmed)
        print(f"  Embedding batch {idx + 1}/{len(batches)} "
              f"({len(trimmed)} docs, ~{est_tokens} tokens)...")
        result = voyage_client.embed(trimmed, model=VOYAGE_MODEL, input_type="document")
        all_embs.extend(result.embeddings)
        if idx < len(batches) - 1:
            print(f"  Rate-limit pause ({sleep_secs:.0f}s)...")
            time.sleep(sleep_secs)

    return np.array(all_embs, dtype=np.float32)


def cosine_similarity_matrix(embs: np.ndarray) -> np.ndarray:
    norms = np.linalg.norm(embs, axis=1, keepdims=True)
    embs_n = embs / np.maximum(norms, 1e-9)
    return embs_n @ embs_n.T


# ── Similarity analysis ────────────────────────────────────────────────────────
def find_similar_pairs(sim_matrix: np.ndarray, labels: list[str], top_k: int) -> list[dict]:
    n      = sim_matrix.shape[0]
    pairs  = []
    for i in range(n):
        for j in range(i + 1, n):
            pairs.append({
                "idx_a": i, "idx_b": j,
                "label_a": labels[i], "label_b": labels[j],
                "similarity": float(sim_matrix[i, j]),
            })
    return sorted(pairs, key=lambda x: x["similarity"], reverse=True)[:top_k]


def cluster_embeddings(embs: np.ndarray):
    from sklearn.cluster import DBSCAN
    dist_matrix = 1.0 - cosine_similarity_matrix(embs)
    dist_matrix = np.clip(dist_matrix, 0, None)
    db = DBSCAN(eps=DBSCAN_EPS, min_samples=DBSCAN_MIN, metric="precomputed")
    labels = db.fit_predict(dist_matrix)
    return labels


def tsne_plot(embs: np.ndarray, labels_list: list[str], cluster_labels: np.ndarray,
              title: str, save_path: Path) -> None:
    from sklearn.manifold import TSNE

    n = len(embs)
    perplexity = min(5, n - 1)
    tsne   = TSNE(n_components=2, random_state=42, perplexity=perplexity,
                  metric="cosine", init="pca", learning_rate="auto")
    coords = tsne.fit_transform(embs)

    fig, ax = plt.subplots(figsize=(10, 7))
    unique_clusters = sorted(set(cluster_labels))
    palette = sns.color_palette("tab10", len(unique_clusters))

    for i, (x, y) in enumerate(coords):
        cl = cluster_labels[i]
        color = palette[unique_clusters.index(cl)]
        ax.scatter(x, y, color=color, s=100, zorder=3,
                   edgecolors="white", linewidths=0.5)
        ax.annotate(labels_list[i], (x, y), textcoords="offset points",
                    xytext=(5, 5), fontsize=7, alpha=0.85)

    ax.set_title(title, fontweight="bold")
    ax.set_xlabel("t-SNE 1")
    ax.set_ylabel("t-SNE 2")
    plt.tight_layout()
    plt.savefig(save_path, dpi=150, bbox_inches="tight")
    plt.close()
    print(f"  Saved: {save_path}")


# ── Main ───────────────────────────────────────────────────────────────────────
def main() -> None:
    voyage_key = os.getenv("VOYAGEAI_API_KEY") or os.getenv("VOYAGE_API_KEY")
    if not voyage_key:
        print("ERROR: VOYAGEAI_API_KEY not set in .env. Exiting.")
        return

    try:
        import voyageai
    except ImportError:
        print("ERROR: voyageai not installed. Run: pip install voyageai")
        return

    vc = voyageai.Client(api_key=voyage_key)

    print("Loading data...")
    df = load_data()
    print(f"  {len(df)} cases loaded")

    # Build document corpus: one entry per (case, version)
    texts   = []
    labels  = []
    meta    = []

    for _, row in df.iterrows():
        mrn = int(row["mrn"])
        for ver, text in [("v1", row["v1_text"]), ("v2", row["v2_text"])]:
            if len(text) > 50:
                texts.append(text[:1500])
                label = f"{row['surgeon_last'][:6]}/{row['patient_initials']}-{ver}"
                labels.append(label)
                meta.append({
                    "mrn": mrn,
                    "surgeon": row["surgeon_last"],
                    "patient_initials": row["patient_initials"],
                    "version": ver,
                    "fab_features": row["fab_features"],
                })

    print(f"  Documents to embed: {len(texts)}")

    # ── Embeddings ──────────────────────────────────────────────────────────────
    emb_cache = OUT_DIR / "embeddings.npy"
    label_cache = OUT_DIR / "embedding_labels.json"

    if emb_cache.exists() and label_cache.exists():
        print("  Loading cached embeddings...")
        embs        = np.load(str(emb_cache))
        cached_lbls = json.load(open(label_cache))
        if cached_lbls == labels:
            print("  Cache valid — using cached embeddings")
        else:
            print("  Cache mismatch — re-embedding")
            embs = get_embeddings(texts, vc)
            np.save(str(emb_cache), embs)
            json.dump(labels, open(label_cache, "w"), indent=2)
    else:
        print("  Embedding documents with Voyage AI...")
        embs = get_embeddings(texts, vc)
        np.save(str(emb_cache), embs)
        json.dump(labels, open(label_cache, "w"), indent=2)

    print(f"  Embeddings shape: {embs.shape}")

    # ── Similarity analysis ─────────────────────────────────────────────────────
    sim_matrix = cosine_similarity_matrix(embs)
    top_pairs  = find_similar_pairs(sim_matrix, labels, TOP_K_PAIRS)

    print(f"\n  Top {TOP_K_PAIRS} similar pairs:")
    for p in top_pairs:
        print(f"    {p['label_a']} <-> {p['label_b']}  sim={p['similarity']:.3f}")

    json.dump(top_pairs, open(OUT_DIR / "similar_case_pairs.json", "w"), indent=2)

    # ── v1/v2 within-case similarity ────────────────────────────────────────────
    within_case = []
    for m in meta:
        if m["version"] == "v1":
            mrn    = m["mrn"]
            idx_v1 = labels.index(f"{m['surgeon'][:6]}/{m['patient_initials']}-v1")
            v2_lbl = f"{m['surgeon'][:6]}/{m['patient_initials']}-v2"
            if v2_lbl in labels:
                idx_v2 = labels.index(v2_lbl)
                within_case.append({
                    "mrn":    mrn,
                    "label":  f"{m['surgeon']}/{m['patient_initials']}",
                    "fab_features": m["fab_features"],
                    "v1_v2_similarity": float(sim_matrix[idx_v1, idx_v2]),
                })

    df_within = pd.DataFrame(within_case).sort_values("v1_v2_similarity")
    print(f"\n  v1↔v2 within-case similarity (mean={df_within['v1_v2_similarity'].mean():.3f}):")
    print(df_within[["label", "fab_features", "v1_v2_similarity"]].to_string(index=False))
    df_within.to_csv(OUT_DIR / "within_case_similarity.csv", index=False)

    # ── DBSCAN clustering ────────────────────────────────────────────────────────
    cluster_labels = cluster_embeddings(embs)
    n_clusters     = len(set(cluster_labels)) - (1 if -1 in cluster_labels else 0)
    print(f"\n  DBSCAN: {n_clusters} clusters, {(cluster_labels==-1).sum()} noise points")

    error_clusters = {}
    for ci in sorted(set(cluster_labels)):
        members = [labels[i] for i, c in enumerate(cluster_labels) if c == ci]
        fab_in_cluster = [meta[i]["fab_features"] for i, c in enumerate(cluster_labels)
                          if c == ci]
        error_clusters[str(ci)] = {
            "cluster_id": int(ci),
            "label": "noise" if ci == -1 else f"cluster_{ci}",
            "members": members,
            "fabricated_features": fab_in_cluster,
            "size": len(members),
        }
    json.dump(error_clusters, open(OUT_DIR / "error_clusters.json", "w"), indent=2)

    # ── t-SNE visualisation ──────────────────────────────────────────────────────
    tsne_plot(
        embs, labels, cluster_labels,
        title=f"t-SNE — v1 vs v2 Summary Embeddings  (Voyage AI: {VOYAGE_MODEL})",
        save_path=REPORTS_DIR / "phase2_similarity_map_tsne.png",
    )

    # ── Within-case similarity bar chart ────────────────────────────────────────
    fig, ax = plt.subplots(figsize=(10, 5))
    colors  = ["#e74c3c" if s < 0.8 else "#2ecc71" for s in df_within["v1_v2_similarity"]]
    ax.barh(df_within["label"], df_within["v1_v2_similarity"],
            color=colors, edgecolor="white")
    ax.axvline(0.8, color="black", linestyle="--", alpha=0.5, label="0.8 threshold")
    ax.set_xlabel("Cosine Similarity")
    ax.set_title("v1 ↔ v2 Within-Case Semantic Similarity\n"
                 "(< 0.8 = structurally divergent summaries)", fontweight="bold")
    ax.set_xlim(0, 1)
    ax.legend()
    plt.tight_layout()
    plt.savefig(REPORTS_DIR / "phase2_within_case_similarity.png", dpi=150,
                bbox_inches="tight")
    plt.close()

    # ── Summary report ───────────────────────────────────────────────────────────
    report = {
        "n_documents": len(texts),
        "embedding_model": VOYAGE_MODEL,
        "n_clusters": n_clusters,
        "mean_within_case_similarity": round(
            float(df_within["v1_v2_similarity"].mean()), 4
        ),
        "min_within_case_similarity":  round(
            float(df_within["v1_v2_similarity"].min()), 4
        ),
        "top_similar_pairs": top_pairs[:5],
    }
    json.dump(report, open(OUT_DIR / "similarity_analysis_report.json", "w"), indent=2)

    print(f"\nPhase 2 complete. Outputs in: {OUT_DIR}")
    print(f"Figures in: {REPORTS_DIR}")


if __name__ == "__main__":
    main()
