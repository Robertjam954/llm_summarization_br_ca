"""
Build comprehensive Word report for the V1/V2 fabrication analysis.
Output: reports/prompt_v1_cs4_v2_gpt5.2/fabrication_analysis_report.docx
"""
import json, sys
from pathlib import Path
from datetime import date
import pandas as pd

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ─────────────────────────────────────────────────────────────────────
PROJECT_ROOT = Path(r"C:\Users\jamesr4\OneDrive - Memorial Sloan Kettering Cancer Center"
                    r"\Documents\GitHub\llm_summarization_br_ca")
DATA_PRIVATE = Path(r"C:\Users\jamesr4\loc\data_private")

REPORT_DIR   = PROJECT_ROOT / "reports" / "prompt_v1_cs4_v2_gpt5.2"
DEEPEVAL_OUT = PROJECT_ROOT / "data" / "processed" / "deepeval_runs" / "fab_v1_v2"
FEAT_OUT     = PROJECT_ROOT / "data" / "processed" / "source_doc_features_v3_simple"
CACHE_FILE   = PROJECT_ROOT / "fabrication_analysis" / "cache" / "llm_judge_cache.json"
OUT_DOCX     = REPORT_DIR / "fabrication_analysis_report.docx"

# ── Figures (ordered for report) ──────────────────────────────────────────────
FIGS = [
    ("fig_verdict",   REPORT_DIR / "fab_human_gt_v2_verdict.png",
     "Figure 1. V2 verdict for all 14 human-confirmed V1 fabrications (bar + pie)."),
    ("fig_heatmap_gt", REPORT_DIR / "fab_human_gt_v2_heatmap.png",
     "Figure 2. Side-by-side heatmap: V1 human annotation vs V2 LLM-judge verdict per case."),
    ("fig_pipeline",  REPORT_DIR / "fab_v1_v2_pipeline_summary.png",
     "Figure 3. Pipeline summary: correction pie, per-feature fabrication rate, overall rate."),
    ("fig_heatmap",   REPORT_DIR / "fab_v1_v2_heatmap.png",
     "Figure 4. Fabrication detection heatmap — LLM judge V1 vs V2 (confirmed-fab features)."),
    ("fig_annot",     REPORT_DIR / "fab_v1_v2_annotation_distribution.png",
     "Figure 5. LLM judge verdict distribution for confirmed-fab features (V1 vs V2)."),
]

# ── Load data ─────────────────────────────────────────────────────────────────
human_gt  = pd.read_csv(DEEPEVAL_OUT / "human_gt_v2_correction.csv")

def _code_label(c):
    m = {1: "CORRECT", 2: "OMISSION", 3: "FABRICATION", "N/A": "N/A"}
    try:
        return m.get(int(c), m.get(str(c), str(c)))
    except (TypeError, ValueError):
        return m.get(str(c), str(c))

human_gt["v2_label"] = human_gt["v2_llm_code"].apply(_code_label)
llm_comp  = pd.read_csv(DEEPEVAL_OUT / "fab_correction_comparison.csv")
raw_val   = pd.read_csv(DEEPEVAL_OUT / "fab_v1_v2_validation_raw.csv")
cache     = json.loads(CACHE_FILE.read_text()) if CACHE_FILE.exists() else {}

feat_jsons = sorted(FEAT_OUT.glob("fab_case_*.json"))
feat_records = []
for f in feat_jsons:
    d = json.loads(f.read_text())
    feat_records.append({
        "mrn":            d["mrn"],
        "case_folder":    d["case_folder"],
        "fab_features":   ", ".join(d.get("fab_features", [])),
        "n_docs_cached":  d["n_docs_cached"],
        "char_count":     d["text_features"]["char_count"],
        "word_count":     d["text_features"]["word_count"],
        "lexical_diversity": round(d["text_features"]["lexical_diversity"], 3),
    })
feat_df = pd.DataFrame(feat_records)

# ── Helpers ───────────────────────────────────────────────────────────────────
def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def add_table(doc, df, col_widths=None, header_color="1F4E79"):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(8)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), header_color)
        tcPr.append(shd)

    # Data rows
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val) if val is not None else ""
            cells[i].paragraphs[0].runs[0].font.size = Pt(7.5)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)

    return table

def add_figure(doc, path, caption, width=6.0):
    if Path(path).exists():
        doc.add_picture(str(path), width=Inches(width))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size = Pt(8.5)
        cp.runs[0].italic = True
    else:
        doc.add_paragraph(f"[Figure not found: {path}]")

# ── Build document ─────────────────────────────────────────────────────────────
doc = Document()

# --- Page margins
from docx.oxml import OxmlElement
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1.0)
section.top_margin  = section.bottom_margin = Inches(1.0)

# ── Title ─────────────────────────────────────────────────────────────────────
t = doc.add_heading("Fabrication Analysis Report", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph(
    "V1 Prompt (Claude Sonnet 4 / narrative) vs V2 Prompt (GPT-4o / structured JSON)\n"
    "14 Human-Confirmed AI Fabrication Cases"
)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(11)

meta = doc.add_paragraph(
    f"Generated: {date.today().isoformat()}   |   "
    f"Analysis: prompt_v1_cs4_v2_gpt5.2   |   "
    f"LLM Judge: GPT-4o"
)
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(9)
meta.runs[0].italic = True
doc.add_page_break()

# ── 1. Executive Summary ──────────────────────────────────────────────────────
heading(doc, "1. Executive Summary")

n_total   = len(human_gt)
n_correct = int(human_gt["v2_correct"].sum())
n_omit    = int(human_gt["v2_omission"].sum())
n_v2fab   = int(human_gt["v2_fabrication"].sum())
n_fixed   = int(human_gt["corrected"].sum())

summary_text = (
    f"This report presents the results of an automated LLM-based evaluation comparing "
    f"V1 (narrative, Claude Sonnet 4 prompt) and V2 (structured JSON, GPT-4o prompt) "
    f"AI-generated oncology summaries across 14 cases with human-confirmed fabrications.\n\n"
    f"Human reviewers annotated {n_total} feature-level fabrications (status_ai = 3) across "
    f"14 breast oncology cases in the validation datasheet. The V2 prompt was then assessed "
    f"by a GPT-4o LLM judge reading the full structured output.\n\n"
    f"Key finding: V2 eliminated {n_fixed}/14 fabrications ({100*n_fixed/n_total:.0f}%), "
    f"with {n_correct} correctly reported, {n_omit} omitted (but not fabricated), and "
    f"only {n_v2fab} remaining fabrications."
)
doc.add_paragraph(summary_text)

doc.add_paragraph()
heading(doc, "Key Metrics", level=2)

metrics_summary = pd.DataFrame([
    ["Human-confirmed V1 fabrications",    f"{n_total}/14 cases"],
    ["V2 CORRECT (code 1)",                f"{n_correct}/14  ({100*n_correct/n_total:.0f}%)"],
    ["V2 OMISSION (code 2)",               f"{n_omit}/14  ({100*n_omit/n_total:.0f}%)"],
    ["V2 still FABRICATING (code 3)",      f"{n_v2fab}/14  ({100*n_v2fab/n_total:.0f}%)"],
    ["Not fabricated in V2 (codes 1+2+NA)", f"{n_fixed}/14  ({100*n_fixed/n_total:.0f}%)"],
    ["LLM judge model",                    "GPT-4o (gpt-4o)"],
    ["V1 prompt style",                    "Narrative (Claude Sonnet 4)"],
    ["V2 prompt style",                    "Structured JSON (GPT-4o)"],
], columns=["Metric", "Value"])
add_table(doc, metrics_summary, col_widths=[3.5, 2.5])
doc.add_page_break()

# ── 2. Human GT vs V2 Correction Results ─────────────────────────────────────
heading(doc, "2. Human GT → V2 Fabrication Correction (Primary Analysis)")
doc.add_paragraph(
    "For each of the 14 human-confirmed V1 fabrications (status_ai = 3 in the identified "
    "validation sheet), the V2 structured JSON summary was evaluated by GPT-4o. "
    "The judge read the full flattened V2 JSON (including nested lesion features) and "
    "the first 3000 characters of source OCR text."
)

display_cols = ["mrn", "case_folder", "feature_display",
                "v1_human_code", "v2_llm_code", "v2_label", "corrected", "v2_confidence"]
display_df = human_gt[display_cols].copy()
display_df.columns = ["MRN", "Case", "Feature", "V1 (Human)", "V2 Code",
                      "V2 Verdict", "Corrected?", "Confidence"]
display_df["Corrected?"] = display_df["Corrected?"].map({True: "YES", False: "NO"})
display_df["Confidence"] = display_df["Confidence"].map(lambda x: f"{float(x):.2f}")
add_table(doc, display_df, col_widths=[1.0, 1.1, 1.6, 0.65, 0.65, 0.9, 0.75, 0.75])

doc.add_paragraph()
heading(doc, "V2 Reasoning (LLM Judge)", level=2)
doc.add_paragraph("Reasoning provided by GPT-4o for each case:")

for _, row in human_gt.iterrows():
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{row['case_folder']} | {row['feature_display']} → {row['v2_label']}: ")
    run.bold = True
    run.font.size = Pt(8.5)
    reason_run = p.add_run(str(row.get("v2_reasoning", ""))[:400])
    reason_run.font.size = Pt(8)
    reason_run.italic = True

doc.add_page_break()

# ── 3. Figures ────────────────────────────────────────────────────────────────
heading(doc, "3. Figures")

for fig_id, fig_path, caption in FIGS:
    add_figure(doc, fig_path, caption, width=5.8)
    doc.add_paragraph()

doc.add_page_break()

# ── 4. LLM Judge Comparison (Notebook 02) ────────────────────────────────────
heading(doc, "4. LLM Judge Comparison — Both Prompts (Notebook 02 Analysis)")
doc.add_paragraph(
    "This table shows LLM judge verdicts for V1 and V2 when BOTH were evaluated "
    "by GPT-4o on the confirmed-fabrication features. Note: V1 text was the narrative DOCX; "
    "V2 text was the extracted feature value (may underestimate V2 accuracy vs the primary analysis)."
)
comp_display = llm_comp[["mrn", "case_folder", "feature", "element_display",
                           "v1_annotation", "v2_annotation",
                           "v1_fabrication", "v2_fabrication", "corrected_in_v2"]].copy()
comp_display.columns = ["MRN", "Case", "Feature", "Element",
                         "V1 Code", "V2 Code",
                         "V1 Fab?", "V2 Fab?", "Corrected?"]
comp_display["V1 Fab?"] = comp_display["V1 Fab?"].map({True: "YES", False: "NO"})
comp_display["V2 Fab?"] = comp_display["V2 Fab?"].map({True: "YES", False: "NO"})
comp_display["Corrected?"] = comp_display["Corrected?"].map({True: "YES", False: "NO"})
add_table(doc, comp_display, col_widths=[0.9, 1.0, 1.0, 1.4, 0.6, 0.6, 0.6, 0.6, 0.8])
doc.add_page_break()

# ── 5. Source Document Features ───────────────────────────────────────────────
heading(doc, "5. Source Document Features (Phase 1 — OCR Text Extraction)")
doc.add_paragraph(
    f"Per-case OCR text statistics extracted from {len(feat_jsons)} cached source documents "
    f"(PDF/DOCX). Output: data/processed/source_doc_features_v3_simple/"
)
feat_disp = feat_df.copy()
feat_disp.columns = ["MRN", "Case", "Fab Feature(s)", "Docs Cached",
                      "Char Count", "Word Count", "Lex Diversity"]
add_table(doc, feat_disp, col_widths=[0.9, 1.1, 1.5, 0.75, 0.85, 0.85, 0.85])

doc.add_paragraph()
heading(doc, "JSON Output Files", level=2)
for f in sorted(feat_jsons):
    d = json.loads(f.read_text())
    doc.add_paragraph(
        f"{f.name}  |  MRN: {d['mrn']}  |  Case: {d['case_folder']}  |  "
        f"Fab: {', '.join(d.get('fab_features', []))}  |  "
        f"{d['n_docs_cached']} docs  {d['text_features']['char_count']:,} chars",
        style="List Bullet"
    ).runs[0].font.size = Pt(8)
doc.add_page_break()

# ── 6. Output File Inventory ──────────────────────────────────────────────────
heading(doc, "6. Output File Inventory")

all_outputs = []
for path in sorted((DEEPEVAL_OUT).glob("*")):
    all_outputs.append([path.name, path.suffix.upper().lstrip("."),
                         f"{path.stat().st_size:,}", str(path)])
for path in sorted(REPORT_DIR.glob("*.png")):
    all_outputs.append([path.name, "PNG",
                         f"{path.stat().st_size:,}", str(path)])
for path in sorted(FEAT_OUT.glob("fab_case_*.json")):
    all_outputs.append([path.name, "JSON",
                         f"{path.stat().st_size:,}", str(path)])
if CACHE_FILE.exists():
    p = CACHE_FILE
    all_outputs.append([p.name, "JSON", f"{p.stat().st_size:,}", str(p)])

inv_df = pd.DataFrame(all_outputs, columns=["File", "Type", "Size (bytes)", "Full Path"])
add_table(doc, inv_df, col_widths=[1.8, 0.5, 0.8, 3.4])

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT_DOCX)
print(f"Report saved: {OUT_DOCX}")
print(f"Size: {OUT_DOCX.stat().st_size:,} bytes")
